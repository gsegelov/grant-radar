"""
tools/export_tools.py

Export functions for pipeline output assembly.
Called directly by app.py after all agents complete — these are NOT tools
attached to any agent. They are plain Python functions that take pipeline
results and return downloadable file bytes.
"""

import csv
import io
from docx import Document
from docx.shared import Pt, RGBColor
from models.schemas import ScoredGrant, GrantBrief, ProposalDraft, ComplianceReport, OrgProfile


def build_csv(scored_grants: list[ScoredGrant]) -> bytes:
    """
    Build a ranked grant opportunity table as CSV bytes.
    Called by app.py after Phase 5 completes.
    Returns bytes so Gradio can serve it as a direct download.
    """
    output = io.StringIO()              # in-memory text buffer — no file written to disk
    writer = csv.writer(output)

    # header row
    writer.writerow([
        "Rank", "Grant Name", "Funder", "Score",
        "Recommended", "Disqualified", "Award Range",
        "Deadline", "Rationale", "Top Alignment Points", "Gaps"
    ])

    # sort by score descending so rank 1 = best fit
    sorted_grants = sorted(scored_grants, key=lambda g: g.score, reverse=True)

    for rank, sg in enumerate(sorted_grants, start=1):
        writer.writerow([
            rank,
            sg.grant.candidate.name,        # ScoredGrant → GrantData → GrantCandidate
            sg.grant.candidate.funder,
            sg.score,
            sg.recommended,
            sg.disqualified,
            f"${sg.grant.award_min}–${sg.grant.award_max}",
            sg.grant.deadline,
            sg.rationale,
            "; ".join(sg.top_alignment_points),     # list → semicolon-separated string
            "; ".join(sg.gaps)
        ])

    return output.getvalue().encode("utf-8")    # convert string to bytes for download


def build_docx(
    org_profile: OrgProfile,
    scored_grants: list[ScoredGrant],
    briefs: list[GrantBrief],
    drafts: list[ProposalDraft],
    reports: list[ComplianceReport]
) -> bytes:
    """
    Assemble the full output document as DOCX bytes.
    One chapter per grant — Brief, Draft, Compliance Report in sequence.
    Called by app.py after all pipeline phases complete.
    """
    doc = Document()                    # blank Word document in memory

    # ── COVER PAGE ────────────────────────────────────────────────────────
    doc.add_heading("GrantRadar Report", 0)
    doc.add_heading(f"Organization: {org_profile.name}", 1)
    doc.add_paragraph(f"Mission: {org_profile.mission_summary}")
    doc.add_paragraph(f"Sector: {org_profile.primary_sector} | Geography: {org_profile.geographic_scope}")
    doc.add_paragraph(f"Grant Size Target: {org_profile.grant_size_target}")
    doc.add_page_break()

    # ── SCORED GRANT SUMMARY TABLE ────────────────────────────────────────
    doc.add_heading("Grant Opportunity Summary", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    # header row
    headers = ["Rank", "Grant Name", "Funder", "Score", "Recommended"]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    # data rows sorted by score
    sorted_grants = sorted(scored_grants, key=lambda g: g.score, reverse=True)
    for rank, sg in enumerate(sorted_grants, start=1):
        row = table.add_row().cells
        row[0].text = str(rank)
        row[1].text = sg.grant.candidate.name
        row[2].text = sg.grant.candidate.funder
        row[3].text = str(sg.score)
        row[4].text = "Yes" if sg.recommended else "No"

    doc.add_page_break()

    # ── ONE CHAPTER PER GRANT ─────────────────────────────────────────────
    # lookup dicts so we can find brief/draft/report by grant name
    brief_lookup = {b.grant_name: b for b in briefs}
    draft_lookup = {d.grant_name: d for d in drafts}
    report_lookup = {r.grant_name: r for r in reports}

    for sg in sorted_grants:
        # skip disqualified grants from chapter section — they still show in the summary table
        if sg.disqualified:
            continue
        grant_name = sg.grant.candidate.name

        # ── SECTION 1: STRATEGIC BRIEF ────────────────────────────────────
        doc.add_heading(grant_name, 1)
        doc.add_heading("Strategic Brief", 2)

        brief = brief_lookup.get(grant_name)
        if brief:
            doc.add_paragraph(f"Recommended Framing: {brief.recommended_framing}")
            doc.add_paragraph("Priorities to Emphasize:")
            for point in brief.priorities_to_emphasize:
                doc.add_paragraph(point, style="List Bullet")
            doc.add_paragraph(f"Narrative Hook: {brief.narrative_hook}")
        else:
            doc.add_paragraph("Strategic brief not available for this grant.")

        # ── SECTION 2: PROPOSAL DRAFT ─────────────────────────────────────
        doc.add_heading("Proposal Draft", 2)

        draft = draft_lookup.get(grant_name)
        if draft:
            if draft.needs_manual_completion:
                # red warning banner if quality guardrail tripped twice
                warning = doc.add_paragraph("⚠ This draft requires manual completion.")
                warning.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

            doc.add_heading("Organizational Capacity", 3)
            doc.add_paragraph(draft.organizational_capacity)

            doc.add_heading("Problem Statement", 3)
            doc.add_paragraph(draft.problem_statement)

            doc.add_heading("Project Description", 3)
            doc.add_paragraph(draft.project_description)

            doc.add_heading("Evaluation Plan", 3)
            doc.add_paragraph(draft.evaluation_plan)
        else:
            doc.add_paragraph("Proposal draft not available for this grant.")

        # ── SECTION 3: COMPLIANCE REPORT ──────────────────────────────────
        doc.add_heading("Compliance Report", 2)

        report = report_lookup.get(grant_name)
        if report:
            doc.add_paragraph(f"Overall Readiness: {report.overall_readiness}")

            doc.add_paragraph("Requirements Addressed:")
            for item in report.requirements_addressed:
                doc.add_paragraph(item, style="List Bullet")

            doc.add_paragraph("Requirements Missing:")
            for item in report.requirements_missing:
                doc.add_paragraph(item, style="List Bullet")

            doc.add_paragraph("Priority Fixes:")
            for fix in report.priority_fixes:
                doc.add_paragraph(fix, style="List Bullet")
        else:
            doc.add_paragraph("Compliance report not available for this grant.")

        doc.add_page_break()

    # ── SAVE TO BYTES ─────────────────────────────────────────────────────
    buffer = io.BytesIO()           # in-memory bytes buffer
    doc.save(buffer)                # save the Word document into the buffer
    buffer.seek(0)                  # move cursor back to start so it can be read
    return buffer.read()            # return raw bytes for Gradio download


if __name__ == "__main__":
    from models.schemas import GrantCandidate, GrantData

    # dummy objects to test export functions
    candidate = GrantCandidate(name="Test Grant", funder="Test Foundation", url="https://example.com", description="A test grant", source_query="test query")
    grant_data = GrantData(candidate=candidate)
    scored = ScoredGrant(grant=grant_data, score=85, rationale="Good fit", top_alignment_points=["Point 1"], recommended=True)
    brief = GrantBrief(grant_name="Test Grant", recommended_framing="Lead with impact", priorities_to_emphasize=["Education"], narrative_hook="Opening hook here")
    draft = ProposalDraft(grant_name="Test Grant", organizational_capacity="We have strong capacity", problem_statement="The problem is X", project_description="We will do Y", evaluation_plan="We will measure Z")
    report = ComplianceReport(grant_name="Test Grant", overall_readiness="Strong", priority_fixes=["Fix 1", "Fix 2", "Fix 3"])
    profile = OrgProfile(name="Test Org", mission_summary="We help people", primary_sector="Education", eligibility_tags=["501c3"], search_queries=["test query"])

    # test CSV
    csv_bytes = build_csv([scored])
    print("CSV OK —", len(csv_bytes), "bytes")

    # test DOCX
    docx_bytes = build_docx(profile, [scored], [brief], [draft], [report])
    with open("test_output.docx", "wb") as f:
        f.write(docx_bytes)
    print("DOCX OK — saved as test_output.docx")